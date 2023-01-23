import os
import pandas as pd
from os import walk
from os.path import dirname, abspath
script_path = abspath(dirname(__file__))

# need to pip install python-docx 
import docx
import io
import zipfile

# get the title of the folder

# walk through the directory
current_directory = os.getcwd()
on_mac = True

if(current_directory.startswith("C:\\")):
    print("YOU ARE ON WINDOWS")
    on_mac = False

if on_mac:

    # handle zips
    zip_files = []

    # find all .zip files, extract them,
    # and convert them to folders
    for files_and_folders in os.walk(current_directory):
        for file_name in files_and_folders[2]:
            if file_name.endswith('.zip'):

                # check if it exists already
                if not os.path.isdir(current_directory + '/' + file_name[:-4]):
                    os.mkdir(current_directory + '/' + file_name[:-4])

                with zipfile.ZipFile(file_name, "r") as z:
                    z.extractall(current_directory + '/' + file_name[:-4])

    # for each folder in the current directory
    for files_and_folders in os.walk(current_directory):
        # if there's a folder
        if(files_and_folders[0].startswith('./')):
            folder_name = files_and_folders[0]

            if(files_and_folders[2]):
                first_file = True
                for file_in_folder in files_and_folders[2]:
                    try:
                        # folder contains .docx files
                        if file_in_folder.endswith('.docx'):
                            
                            # the .docx file
                            content = '\n'.join([p.text for p in docx.Document(folder_name + "/" + file_in_folder).paragraphs])
                            df = pd.read_csv(io.StringIO(content), encoding='latin1', sep=',', on_bad_lines='skip')

                            # check if the column starts with the windows quotes like in
                            # all of the sample CSV-type docs
                            if df.columns[0].startswith('“') and df.columns[0].endswith('”'):
                                
                                if first_file:
                                    new_folder_path = current_directory + "/" + files_and_folders[0][2:] + "_CSVs"

                                    # folder of foldername_CSVS/ created in
                                    # the same directory as the script
                                    # this is where we put the CSVS
                                    
                                    # check if it exists already
                                    if not os.path.isdir(new_folder_path):
                                        os.mkdir(new_folder_path)

                                    first_file = False

                                csv_filename = file_in_folder[:-4] + "csv"

                                print("printing csv {}".format(new_folder_path + "/" + csv_filename))
                                df.to_csv(new_folder_path + "/" + csv_filename, index=False)
                            
                            else:
                                print(df.columns[0])

                        # folder contains .docx files
                        elif file_in_folder.endswith('.txt'):
                            df = pd.read_csv(file_in_folder, encoding='latin1', sep=',', on_bad_lines='skip')
                                
                            if first_file:
                                new_folder_path = current_directory + "/" + files_and_folders[0][2:] + "_CSVs"

                                # folder of foldername_CSVS/ created in
                                # the same directory as the script
                                # this is where we put the CSVS
                                
                                # check if it exists already
                                if not os.path.isdir(new_folder_path):
                                    os.mkdir(new_folder_path)

                                first_file = False

                            csv_filename = file_in_folder[:-3] + "csv"

                            print("printing csv {}".format(new_folder_path + "/" + csv_filename))
                            df.to_csv(new_folder_path + "/" + csv_filename, index=False)

                                
                    except Exception as e: 
                        print(e)
                        continue

# if running on Windows
else:
    # handle zips
    zip_files = []

    # find all .zip files, extract them,
    # and convert them to folders
    for files_and_folders in os.walk(current_directory):
        for file_name in files_and_folders[2]:
            if file_name.endswith('.zip'):

                # check if it exists already
                if not os.path.isdir(current_directory + '\\' + file_name[:-4]):
                    os.mkdir(current_directory + '\\' + file_name[:-4])

                with zipfile.ZipFile(file_name, "r") as z:
                    z.extractall(current_directory + '\\' + file_name[:-4])


    # for each folder in the current directory
    for files_and_folders in os.walk(current_directory):
        print(files_and_folders)
        # if there's a folder found
        if(files_and_folders[0].startswith(current_directory + '\\')):
            folder_name = files_and_folders[0]

            if(files_and_folders[2]):
                first_file = True
                for file_in_folder in files_and_folders[2]:
                    try:
                        # folder contains .docx files
                        if file_in_folder.endswith('.docx'):
                            
                            # the .docx file
                            content = '\n'.join([p.text for p in docx.Document(folder_name + "\\" + file_in_folder).paragraphs])

                            # some files had encoding issues, these are skipped
                            df = pd.read_csv(io.StringIO(content), encoding='latin1', sep=',', on_bad_lines='skip')

                            # check if the column starts with the windows quotes like in
                            # all of the sample CSV-type docs
                            if df.columns[0].startswith('“') and df.columns[0].endswith('”'):
                                
                                if first_file:
                                    new_folder_path = folder_name + "_CSVs"

                                    # folder of foldername_CSVS/ created in
                                    # the same directory as the script
                                    # this is where we put the CSVS
                                    
                                    # check if it exists already
                                    if not os.path.isdir(new_folder_path):
                                        os.mkdir(new_folder_path)

                                    first_file = False

                                csv_filename = file_in_folder[:-4] + "csv"

                                df.to_csv(new_folder_path + "\\" + csv_filename, index=False)
                            
                            else:
                                print(df.columns[0])

                        # folder contains .txt files
                        elif file_in_folder.endswith('.txt'):
                            path = folder_name + "\\" + file_in_folder

                            # some files have strange encodings, so just skipping for now
                            df = pd.read_csv(path, encoding='latin1', sep=',', on_bad_lines='skip')

                            # for .txt files it's easier and we can just read them
                            if first_file:
                                new_folder_path = folder_name + "_CSVs"

                                # folder of foldername_CSVS/ created in
                                # the same directory as the script
                                # this is where we put the CSVS
                                    
                                # check if it exists already
                                if not os.path.isdir(new_folder_path):
                                    os.mkdir(new_folder_path)

                                    first_file = False

                            csv_filename = file_in_folder[:-3] + "csv"

                            df.to_csv(new_folder_path + "\\" + csv_filename, index=False)

                        else:
                            print(df.columns[0])

                                
                    except Exception as e: 
                        print(e)
                        continue